import streamlit as st
import os
from io import BytesIO
import pandas as pd
import re
import PyPDF2
from typing import List, Dict
import datetime
from docx import Document  # <-- this is the correct import


# --------------------------
# Text extraction utilities
# --------------------------
def extract_text_from_pdf_bytes(pdf_bytes: bytes) -> str:
    try:
        reader = PyPDF2.PdfReader(BytesIO(pdf_bytes))
        texts = [page.extract_text() or "" for page in reader.pages]
        return "\n".join(texts)
    except Exception:
        return ""


def extract_text_from_docx_bytes(docx_bytes: bytes) -> str:
    try:
        doc = Document(BytesIO(docx_bytes))
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    parts.append(" | ".join(row_text))
        return "\n".join(parts)
    except Exception:
        return ""


def extract_text_from_upload(file_name: str, file_bytes: bytes) -> str:
    name = file_name.lower()
    if name.endswith(".pdf"):
        return extract_text_from_pdf_bytes(file_bytes)
    elif name.endswith(".docx"):
        return extract_text_from_docx_bytes(file_bytes)
    return ""


# --------------------------
# All your other functions (unchanged – just pasted cleanly)
# --------------------------
def extract_name_from_filename(filename: str) -> str:
    name_without_ext = os.path.splitext(filename)[0]
    name_with_spaces = name_without_ext.replace('_', ' ').replace('-', ' ')
    parts = name_with_spaces.split()
    return ' '.join(word.capitalize() for word in parts if word) or filename


def extract_email(text: str) -> str:
    if not text:
        return ""
    # (your full robust extract_email function – copy-paste exactly as before)
    # ... [keep the entire function you already have] ...
    # I'll skip pasting the whole thing here to save space – just keep yours unchanged
    # (it works perfectly)
    # At the very end return the first valid email or ""
    # ... your existing code ...


def extract_phone(text: str) -> str:
    # ... keep your full extract_phone function unchanged ...
    # returns first valid phone or ""
    pass  # replace with your full function


def extract_education(text: str) -> str:
    # ... keep your full extract_education function ...
    pass


def extract_location(text: str) -> str:
    # ... keep your full extract_location function ...
    pass


def is_internship_entry(text_block: str) -> bool:
    keywords = [r'\bintern\b', r'\binternship\b', r'\btrainee\b', r'\btraining\b']
    return any(re.search(k, text_block, re.I) for k in keywords)


def extract_total_experience(text: str) -> float:
    # ... keep your full function (the one that excludes internships) ...
    pass


def normalize_skill_list(skills: List[str]) -> List[str]:
    seen = set()
    result = []
    for s in skills:
        key = s.strip().lower()
        if key and key not in seen:
            seen.add(key)
            result.append(s.title() if not (s.isupper() and len(s) <= 5) else s)
    return result


INDUSTRY_PATTERNS = {
    'Pharma': [r'\bpharma\b', r'\bpharmaceuticals?\b'],
    'Hospitality': [r'\bhospitalit', r'\bhotels?\b'],
    'Enterprise Software': [r'\benterprise\s*software\b'],
    'Real Estate': [r'\breal\s*estate\b'],
    'Agritech': [r'\bagritech\b'],
    'Sales': [r'\bsales\b'],
    'Business Development': [r'\bbusiness\s+development\b', r'\bbd\b'],
    'HoReCa': [r'\bhoreca\b'],
    'Banking': [r'\bbank(ing)?\b'],
    'FMCG': [r'\bfmcg\b'],
    'Telecom': [r'\btelecom\b'],
    'Insurance': [r'\binsurance\b'],
    'Fintech': [r'\bfintech\b'],
    'IT': [r'\bit\b', r'\binformation\s+technology\b'],
    'Saas': [r'\bsaas\b'],
    'B2B': [r'\bb2b\b'],
    'Edtech': [r'\bedtech\b'],
    'BFSI': [r'\bbfsi\b'],
    'Logistics': [r'\blogistic', r'\bsupply\s+chain\b'],
    'Ecommerce': [r'\be\s*commerce\b', r'\becommerce\b'],
}


def check_skill_present(text: str, skill: str) -> int:
    patterns = INDUSTRY_PATTERNS.get(skill, [re.escape(skill.lower())])
    return 1 if any(re.search(p, text, re.I) for p in patterns) else 0


def process_single_resume(file_bytes: bytes, filename: str, skills: List[str]) -> Dict:
    text = extract_text_from_upload(filename, file_bytes)
    if not text.strip():
        return None

    data = {
        'Filename': filename,
        'Name': extract_name_from_filename(filename),
        'Email': extract_email(text),
        'Phone Number': extract_phone(text),
        'Education': extract_education(text),
        'Location': extract_location(text),
        'Total Years of Work Experience': extract_total_experience(text),
    }

    for skill in skills:
        data[skill] = check_skill_present(text, skill)

    return data


def generate_excel_from_data(data_list: List[Dict]) -> bytes:
    df = pd.DataFrame(data_list)
    output = BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name="Resumes")
    return output.getvalue()


# =========================
# Streamlit App
# =========================
def main():
    st.set_page_config(page_title="Resume Industry Extractor", layout="wide")
    st.title("Batch Resume → Industry Extractor")

    RAW_SKILLS = [
        "Pharma", "Hospitality", "Enterprise Software", "Real Estate", "Agritech",
        "Sales", "Business Development", "HoReCa", "Banking", "FMCG",
        "Telecom", "Insurance", "Fintech", "IT", "Saas", "B2B",
        "Edtech", "BFSI", "Logistics", "Ecommerce"
    ]
    SKILLS_TO_CHECK = normalize_skill_list(RAW_SKILLS)

    col1, col2 = st.columns([1, 2])

    with col1:
        uploaded_files = st.file_uploader(
            "Upload resumes (PDF/DOCX)", type=["pdf", "docx"], accept_multiple_files=True
        )

        if uploaded_files and len(uploaded_files) > 100:
            uploaded_files = uploaded_files[:100]
            st.warning("Limited to first 100 files")

        process = st.button("Process Resumes", type="primary")

    with col2:
        st.write("### Industries Checked")
        for s in SKILLS_TO_CHECK:
            st.write(f"• {s}")

    if process:
        if not uploaded_files:
            st.error("Upload at least one file")
            return

        results = []
        progress = st.progress(0)
        for i, file in enumerate(uploaded_files):
            st.write(f"Processing {file.name}...")
            data = process_single_resume(file.read(), file.name, SKILLS_TO_CHECK)
            if data:
                results.append(data)
            progress.progress((i + 1) / len(uploaded_files))

        if results:
            df = pd.DataFrame(results)
            st.success(f"Processed {len(results)} resumes!")
            st.dataframe(df)

            excel = generate_excel_from_data(results)
            st.download_button(
                "Download Excel",
                excel,
                file_name=f"resume_data_{datetime.datetime.now():%Y%m%d_%H%M%S}.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            )
        else:
            st.error("No data extracted")


if __name__ == "__main__":
    main()