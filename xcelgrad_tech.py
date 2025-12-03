import streamlit as st
from io import BytesIO
import pandas as pd
import re
import PyPDF2
from typing import List, Dict
from docx import Document

# --------------------------
# Utilities: PDF/DOCX -> text
# --------------------------
def extract_text_from_pdf_bytes(pdf_bytes: bytes) -> str:
    """Extract all text from PDF file"""
    try:
        reader = PyPDF2.PdfReader(BytesIO(pdf_bytes))
    except Exception:
        return ""
    texts = []
    for page in reader.pages:
        try:
            texts.append(page.extract_text() or "")
        except Exception:
            texts.append("")
    return "\n".join(texts)

def extract_text_from_docx_bytes(docx_bytes: bytes) -> str:
    """Extract text from DOCX (paragraphs + tables)"""
    try:
        doc = Document(BytesIO(docx_bytes))
    except Exception:
        return ""
    parts = []
    # paragraphs
    for p in doc.paragraphs:
        if p.text:
            parts.append(p.text)
    # tables
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()]
            if row_text:
                parts.append(" | ".join(row_text))
    return "\n".join(parts)

def extract_text_from_upload(file_name: str, file_bytes: bytes) -> str:
    """Dispatch extraction based on extension"""
    name = file_name.lower()
    if name.endswith(".pdf"):
        return extract_text_from_pdf_bytes(file_bytes)
    elif name.endswith(".docx"):  # .doc not supported reliably without external deps
        return extract_text_from_docx_bytes(file_bytes)
    else:
        return ""

# --------------------------
# Information extraction
# --------------------------
def extract_name(text: str) -> str:
    lines = [l.strip() for l in text.split('\n') if l.strip()]
    if not lines:
        return ""
    first_line = re.sub(r'^(resume|curriculum vitae|cv)[\s:]*', '', lines[0], flags=re.I)
    words = first_line.split()
    if 2 <= len(words) <= 4 and all(word.replace('.', '').isalpha() for word in words):
        return first_line
    if len(lines) > 1:
        second_line = lines[1]
        words = second_line.split()
        if 2 <= len(words) <= 4 and all(word.replace('.', '').isalpha() for word in words):
            return second_line
    return first_line[:50]

def extract_email(text: str) -> str:
    email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
    matches = re.findall(email_pattern, text)
    return matches[0] if matches else ""

def extract_phone(text: str) -> str:
    phone_patterns = [
        r'\+?\d{1,3}[-.\s]?\(?\d{3,5}\)?[-.\s]?\d{3,5}[-.\s]?\d{4}',  # flexible intl
        r'\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}',
        r'\+?\d{10,12}',
        r'\d{3}[-.\s]\d{3}[-.\s]\d{4}'
    ]
    for pattern in phone_patterns:
        matches = re.findall(pattern, text)
        if matches:
            return matches[0]
    return ""

def extract_education(text: str) -> str:
    text_lower = text.lower()
    education_keywords = ['education', 'academic', 'qualification']
    education_start = -1
    for keyword in education_keywords:
        match = re.search(r'\b' + keyword + r'\b', text_lower)
        if match:
            education_start = match.start()
            break
    if education_start == -1:
        degree_pattern = r'\b(bachelor|master|phd|b\.tech|m\.tech|b\.e|m\.e|bsc|msc|bca|mca|diploma)\b'
        matches = re.findall(degree_pattern, text_lower)
        if matches:
            return matches[0].upper()
        return ""
    education_text = text[education_start:education_start + 500]
    degree_pattern = r'(Bachelor[^,\n]*|Master[^,\n]*|PhD[^,\n]*|B\.Tech[^,\n]*|M\.Tech[^,\n]*|B\.E[^,\n]*|M\.E[^,\n]*|BSc[^,\n]*|MSc[^,\n]*|BCA[^,\n]*|MCA[^,\n]*)'
    degree_match = re.search(degree_pattern, education_text, re.I)
    if degree_match:
        return degree_match.group(1).strip()
    lines = [l.strip() for l in education_text.split('\n') if l.strip()]
    return lines[1] if len(lines) > 1 else (lines[0] if lines else "")

# --------------------------
# Skill / Industry matching
# --------------------------
def normalize_skill_list(skills: List[str]) -> List[str]:
    """Deduplicate case-insensitive while preserving order and trim whitespace."""
    seen = set()
    normalized = []
    for s in skills:
        key = s.strip()
        if not key:
            continue
        lower = key.lower()
        if lower not in seen:
            seen.add(lower)
            if key.isupper() and len(key) <= 5:
                display = key
            else:
                display = key.title()
            normalized.append(display)
    return normalized

INDUSTRY_PATTERNS = {
    'Pharma': [r'\bpharma\b', r'\bpharmaceuticals?\b', r'\bpharmaceutical\b'],
    'Hospitality': [r'\bhospitalit(y|ies)\b', r'\bhotels?\b', r'\bfood\s+and\s+beverage\b', r'\bfnb\b'],
    'Enterprise Software': [r'\benterprise[\s\-]?software\b', r'\benterprise\s+apps?\b', r'\benterprise\s+solutions?\b'],
    'Real Estate': [r'\breal[\s\-]?estate\b', r'\bproperty\s+development\b', r'\bproperty\s+management\b'],
    'Agritech': [r'\bagritech\b', r'\bagri[\s\-]?tech\b', r'\bagriculture\b', r'\bfarming\b'],
    'Sales': [r'\bsales\b', r'\bsales\s+professional\b', r'\bsales\s+executive\b'],
    'Business Development': [r'\bbusiness\s+development\b', r'\bbd\s+manager\b', r'\bbusiness\s+dev\b', r'\bbd\b'],
    'HoReCa': [r'\bhoreca\b', r'\bhotel\s+restaurant\s+cafe\b'],
    'Banking': [r'\bbank(ing)?\b', r'\bfinancial\s+services\b'],
    'FMCG': [r'\bfmcg\b', r'\bfast\s+moving\s+consumer\s+goods\b'],
    'TELECOM': [r'\btelecom\b', r'\btelecommunications?\b', r'\btelecoms?\b'],
    'INSURANCE': [r'\binsurance\b', r'\binsurance\s+industry\b'],
    'FINTECH': [r'\bfintech\b', r'\bfinancial\s+technology\b'],
    'IT': [r'\bit\s+sector\b', r'\binformation\s+technology\b', r'\bit\s+services\b', r'\btechnology\s+company\b'],
    'SAAS': [r'\bsaas\b', r'\bsoftware\s+as\s+a\s+service\b'],
    'B2b': [r'\bb2b\b', r'\bbusiness\-to\-business\b'],
    'Edtech': [r'\bedtech\b', r'\beducation\s+technology\b', r'\beducational\s+technology\b'],
    'BFSI': [r'\bbfsi\b', r'\bbanking\s+finance\s+and\s+insurance\b'],
    'Logistic': [r'\blogistic(s)?\b', r'\bsupply\s+chain\b', r'\blogistics?\b'],
    'ECommerce': [r'\be[\s\-]?commerce\b', r'\becommerce\b', r'\bonline\s+retail\b']
}

def check_skill_present(text: str, skill_display_name: str) -> int:
    """Return 1 if any pattern for the skill matches in text (case-insensitive), else 0."""
    if not text or not skill_display_name:
        return 0
    patterns = INDUSTRY_PATTERNS.get(skill_display_name, [r'\b' + re.escape(skill_display_name.lower()) + r'\b'])
    for pat in patterns:
        if re.search(pat, text, flags=re.I):
            return 1
    return 0

# --------------------------
# Resume processing
# --------------------------
def process_single_resume(file_bytes: bytes, filename: str, skills_to_check: List[str]) -> Dict:
    """Process a single resume and return extracted data. Skills are checked in the entire resume text."""
    full_text = extract_text_from_upload(filename, file_bytes)
    if not full_text.strip():
        return None
    name = extract_name(full_text)
    email = extract_email(full_text)
    phone = extract_phone(full_text)
    education = extract_education(full_text)

    data = {
        'Filename': filename,
        'Name': name,
        'Email': email,
        'Phone Number': phone,
        'Education': education
    }

    for skill in skills_to_check:
        data[skill] = check_skill_present(full_text, skill)

    return data

# --------------------------
# Excel generation
# --------------------------
def generate_excel_from_data(all_data: List[Dict]) -> bytes:
    df = pd.DataFrame(all_data)
    output = BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name="Resume_Data")
    return output.getvalue()

# --------------------------
# Streamlit UI (wrapped in main)
# --------------------------
RAW_SKILLS = [
    "Pharma",
    "Hospitality",
    "Enterprise software",
    "Real Estate",
    "Agritech",
    "SALES",
    "Business Development",
    "HoReCa",
    "Banking",
    "FMCG",
    "TELECOM",
    "INSURANCE",
    "FINTECH",
    "IT",
    "SAAS",
    "B2b",
    "sales",
    "EdTech",
    "BFSI",
    "Logistic",
    "ECommerce"
]

SKILLS_TO_CHECK = normalize_skill_list(RAW_SKILLS)

def main():
    st.header("📄 Batch Resume → Industry/Vertical Extractor (Whole-Resume Matching)")
    st.markdown(
        """
        Upload **multiple resumes (PDF or Word .docx)** and get a **single Excel file** with:
        - Basic information (Name, Email, Phone, Education) - extracted from entire resume  
        - Industry/vertical presence indicators (1 if present, 0 if not) - **searched across the entire resume text**
        """
    )

    col1, col2 = st.columns([1, 2])

    with col1:
        uploaded_files = st.file_uploader(
            "Upload Resumes (PDF or DOCX) — up to 100 files",
            type=["pdf", "docx"],
            accept_multiple_files=True
        )

        if uploaded_files:
            if len(uploaded_files) > 100:
                st.warning(f"⚠️ You uploaded {len(uploaded_files)} files. Processing the first 100.")
                uploaded_files = uploaded_files[:100]
            st.success(f"✅ {len(uploaded_files)} file(s) ready to process")

        st.info(
            "**The app will extract:**\n\n"
            "From entire resume:\n"
            "- Filename\n"
            "- Name\n"
            "- Email\n"
            "- Phone Number\n"
            "- Education\n\n"
            "**From entire resume (not only Experience/Projects):**\n"
            "- Industry/Vertical presence (1/0)\n\n"
            "📎 Supported formats: **PDF**, **Word (.docx)**"
        )

        process_button = st.button("🚀 Process All Resumes", type="primary")

    with col2:
        st.subheader("How it works")
        st.write("1. **Upload** multiple resumes (PDF or DOCX).")
        st.write("2. Click **Process All Resumes**.")
        st.write("3. **Download** the Excel file with all extracted information.")
        st.info("ℹ️ Industry matching is now performed across the entire resume text (no section restriction).")
        st.write("")
        st.write("**Industries / Verticals Checked:**")
        cols = st.columns(3)
        for idx, skill in enumerate(SKILLS_TO_CHECK):
            with cols[idx % 3]:
                st.write(f"  • {skill}")

    if process_button:
        if not uploaded_files:
            st.error("⚠️ Please upload at least one resume first.")
        else:
            all_data = []
            progress_bar = st.progress(0)
            status_text = st.empty()

            for idx, uploaded_file in enumerate(uploaded_files):
                status_text.text(f"Processing {idx + 1}/{len(uploaded_files)}: {uploaded_file.name}")
                try:
                    raw_bytes = uploaded_file.read()
                    data = process_single_resume(raw_bytes, uploaded_file.name, SKILLS_TO_CHECK)
                    if data:
                        all_data.append(data)
                    else:
                        st.warning(f"⚠️ Could not extract text from: {uploaded_file.name} (unsupported/empty/corrupt)")
                except Exception as e:
                    st.error(f"❌ Error processing {uploaded_file.name}: {str(e)}")
                progress_bar.progress((idx + 1) / len(uploaded_files))

            status_text.empty()
            progress_bar.empty()

            if not all_data:
                st.error("❌ Could not extract data from any of the uploaded files.")
            else:
                st.success(f"✅ Successfully processed {len(all_data)} out of {len(uploaded_files)} file(s)!")

                st.subheader("📊 Processing Summary")
                col_a, col_b, col_c = st.columns(3)
                with col_a:
                    st.metric("Total Files Uploaded", len(uploaded_files))
                with col_b:
                    st.metric("Successfully Processed", len(all_data))
                with col_c:
                    st.metric("Failed", len(uploaded_files) - len(all_data))

                st.subheader("📈 Industry/Vertical Statistics (from whole resume)")
                skill_counts = {}
                for skill in SKILLS_TO_CHECK:
                    count = sum(1 for data in all_data if data.get(skill) == 1)
                    skill_counts[skill] = count

                num_cols = 4
                skill_items = list(skill_counts.items())
                for i in range(0, len(skill_items), num_cols):
                    skill_cols = st.columns(num_cols)
                    for j in range(num_cols):
                        if i + j < len(skill_items):
                            skill, count = skill_items[i + j]
                            with skill_cols[j]:
                                percentage = (count / len(all_data) * 100) if all_data else 0
                                st.metric(skill, f"{count}/{len(all_data)}", f"{percentage:.0f}%")

                st.subheader("Complete Data Table")
                df_display = pd.DataFrame(all_data)
                st.dataframe(df_display, use_container_width=True)

                with st.spinner("📝 Generating Excel file..."):
                    excel_bytes = generate_excel_from_data(all_data)

                st.download_button(
                    label="📥 Download Excel File with All Data",
                    data=excel_bytes,
                    file_name="batch_resume_industries_extracted.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    type="primary"
                )

    st.markdown("---")
    st.caption("Have a GOOD DAY!!!")

if __name__ == "__main__":
    main()
